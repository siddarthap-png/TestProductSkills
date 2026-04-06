"""Convert a subset of Markdown to .docx (headings, tables, bullets, **bold**)."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK


def add_runs_with_formatting(paragraph, text: str) -> None:
    """Supports **bold** and *italic* (single asterisk pairs)."""
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) >= 2 and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def is_table_sep(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    return bool(re.match(r"^\|[\s\-:|]+\|$", s))


def parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().split("|")[1:-1]]


def main() -> int:
    if len(sys.argv) < 3:
        print("Usage: md_to_docx.py input.md output.docx", file=sys.stderr)
        return 2
    md_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2])
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.LINE)
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=0)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            rows_data = []
            for bl in block:
                if is_table_sep(bl):
                    continue
                rows_data.append(parse_table_row(bl))
            if rows_data:
                ncols = max(len(r) for r in rows_data)
                for r in rows_data:
                    while len(r) < ncols:
                        r.append("")
                table = doc.add_table(rows=len(rows_data), cols=ncols)
                table.style = "Table Grid"
                for ri, row_cells in enumerate(rows_data):
                    for ci, cell_text in enumerate(row_cells):
                        cell = table.rows[ri].cells[ci]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        add_runs_with_formatting(p, cell_text)
            continue

        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs_with_formatting(p, m.group(2))
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_formatting(p, stripped[2:])
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs_with_formatting(p, stripped)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Wrote {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
